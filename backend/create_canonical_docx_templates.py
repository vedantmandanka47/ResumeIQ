"""Generate canonical docxtpl resume templates (run once from backend/)."""

import sys

if sys.version_info < (3, 10):
    sys.exit("Use Python 3.10+: py -3.12 create_canonical_docx_templates.py")

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

TEMPLATE_DIR = Path(__file__).parent / "templete"


def add_bottom_border(paragraph, color_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)


def _hidden_tag(paragraph, text: str, font_name: str) -> None:
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("FFFFFF")


def create_template(filename: str, theme_color_hex: str, font_name: str = "Calibri") -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    theme_color = RGBColor.from_string(theme_color_hex.replace("#", ""))
    dark_gray = RGBColor.from_string("333333")
    light_gray = RGBColor.from_string("666666")

    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run("{{ name }}")
    run_name.font.name = font_name
    run_name.font.size = Pt(24)
    run_name.font.bold = True
    run_name.font.color.rgb = theme_color

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_contact = p_contact.add_run("{{ contact_line }}")
    run_contact.font.name = font_name
    run_contact.font.size = Pt(10)
    run_contact.font.color.rgb = dark_gray

    p_sum_if = doc.add_paragraph()
    _hidden_tag(p_sum_if, "{% if has_summary %}", font_name)
    p_sum_hdr = doc.add_paragraph()
    run_sum_hdr = p_sum_hdr.add_run("PROFESSIONAL SUMMARY")
    run_sum_hdr.font.name = font_name
    run_sum_hdr.font.size = Pt(13)
    run_sum_hdr.font.bold = True
    run_sum_hdr.font.color.rgb = theme_color
    add_bottom_border(p_sum_hdr, theme_color_hex.replace("#", ""))
    p_sum = doc.add_paragraph()
    run_sum = p_sum.add_run("{{ summary }}")
    run_sum.font.name = font_name
    run_sum.font.size = Pt(10.5)
    run_sum.font.color.rgb = dark_gray
    p_sum_end = doc.add_paragraph()
    _hidden_tag(p_sum_end, "{% endif %}", font_name)

    p_exp_if = doc.add_paragraph()
    _hidden_tag(p_exp_if, "{% if has_experience %}", font_name)
    p_exp_hdr = doc.add_paragraph()
    run_exp_hdr = p_exp_hdr.add_run("PROFESSIONAL EXPERIENCE")
    run_exp_hdr.font.name = font_name
    run_exp_hdr.font.size = Pt(13)
    run_exp_hdr.font.bold = True
    run_exp_hdr.font.color.rgb = theme_color
    add_bottom_border(p_exp_hdr, theme_color_hex.replace("#", ""))
    p_loop_exp = doc.add_paragraph()
    _hidden_tag(p_loop_exp, "{% for item in experience %}", font_name)

    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(5.0)
    table.columns[1].width = Inches(2.0)
    row = table.rows[0]
    p_left = row.cells[0].paragraphs[0]
    run_role = p_left.add_run("{{ item.role }}  |  {{ item.company }}")
    run_role.font.name = font_name
    run_role.font.size = Pt(11)
    run_role.font.bold = True
    p_right = row.cells[1].paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_dates = p_right.add_run("{{ item.start_date }} - {{ item.end_date }}")
    run_dates.font.name = font_name
    run_dates.font.size = Pt(10)
    run_dates.font.italic = True

    p_loc = doc.add_paragraph()
    run_loc = p_loc.add_run("{{ item.location }}")
    run_loc.font.name = font_name
    run_loc.font.size = Pt(9.5)
    run_loc.font.italic = True
    run_loc.font.color.rgb = light_gray

    p_bullet_start = doc.add_paragraph()
    _hidden_tag(p_bullet_start, "{% for bullet in item.description %}", font_name)
    p_bullet = doc.add_paragraph(style="List Bullet")
    run_bullet = p_bullet.add_run("{{ bullet }}")
    run_bullet.font.name = font_name
    run_bullet.font.size = Pt(10)
    p_bullet_end = doc.add_paragraph()
    _hidden_tag(p_bullet_end, "{% endfor %}", font_name)
    p_loop_exp_end = doc.add_paragraph()
    _hidden_tag(p_loop_exp_end, "{% endfor %}", font_name)
    p_exp_end = doc.add_paragraph()
    _hidden_tag(p_exp_end, "{% endif %}", font_name)

    p_proj_if = doc.add_paragraph()
    _hidden_tag(p_proj_if, "{% if has_projects %}", font_name)
    p_proj_hdr = doc.add_paragraph()
    run_proj_hdr = p_proj_hdr.add_run("PROJECTS")
    run_proj_hdr.font.name = font_name
    run_proj_hdr.font.size = Pt(13)
    run_proj_hdr.font.bold = True
    run_proj_hdr.font.color.rgb = theme_color
    add_bottom_border(p_proj_hdr, theme_color_hex.replace("#", ""))
    p_loop_proj = doc.add_paragraph()
    _hidden_tag(p_loop_proj, "{% for item in projects %}", font_name)
    p_proj = doc.add_paragraph()
    run_proj = p_proj.add_run("{{ item.name }}  |  {{ item.technologies }}")
    run_proj.font.name = font_name
    run_proj.font.size = Pt(11)
    run_proj.font.bold = True
    p_proj_bullet_start = doc.add_paragraph()
    _hidden_tag(p_proj_bullet_start, "{% for bullet in item.description %}", font_name)
    p_proj_bullet = doc.add_paragraph(style="List Bullet")
    run_proj_bullet = p_proj_bullet.add_run("{{ bullet }}")
    run_proj_bullet.font.name = font_name
    run_proj_bullet.font.size = Pt(10)
    p_proj_bullet_end = doc.add_paragraph()
    _hidden_tag(p_proj_bullet_end, "{% endfor %}", font_name)
    p_loop_proj_end = doc.add_paragraph()
    _hidden_tag(p_loop_proj_end, "{% endfor %}", font_name)
    p_proj_end = doc.add_paragraph()
    _hidden_tag(p_proj_end, "{% endif %}", font_name)

    p_edu_if = doc.add_paragraph()
    _hidden_tag(p_edu_if, "{% if has_education %}", font_name)
    p_edu_hdr = doc.add_paragraph()
    run_edu_hdr = p_edu_hdr.add_run("EDUCATION")
    run_edu_hdr.font.name = font_name
    run_edu_hdr.font.size = Pt(13)
    run_edu_hdr.font.bold = True
    run_edu_hdr.font.color.rgb = theme_color
    add_bottom_border(p_edu_hdr, theme_color_hex.replace("#", ""))
    p_loop_edu = doc.add_paragraph()
    _hidden_tag(p_loop_edu, "{% for item in education %}", font_name)
    p_edu_line = doc.add_paragraph()
    run_edu_line = p_edu_line.add_run("{{ item.degree }}  |  {{ item.institution }}")
    run_edu_line.font.name = font_name
    run_edu_line.font.size = Pt(11)
    run_edu_line.font.bold = True
    p_edu_meta = doc.add_paragraph()
    run_edu_meta = p_edu_meta.add_run("{{ item.graduation_date }}")
    run_edu_meta.font.name = font_name
    run_edu_meta.font.size = Pt(10)
    run_edu_meta.font.italic = True
    run_gpa = p_edu_meta.add_run("{% if item.gpa %}  |  GPA: {{ item.gpa }}{% endif %}")
    run_gpa.font.name = font_name
    run_gpa.font.size = Pt(10)
    p_loop_edu_end = doc.add_paragraph()
    _hidden_tag(p_loop_edu_end, "{% endfor %}", font_name)
    p_edu_end = doc.add_paragraph()
    _hidden_tag(p_edu_end, "{% endif %}", font_name)

    p_sk_if = doc.add_paragraph()
    _hidden_tag(p_sk_if, "{% if has_skills %}", font_name)
    p_sk_hdr = doc.add_paragraph()
    run_sk_hdr = p_sk_hdr.add_run("SKILLS")
    run_sk_hdr.font.name = font_name
    run_sk_hdr.font.size = Pt(13)
    run_sk_hdr.font.bold = True
    run_sk_hdr.font.color.rgb = theme_color
    add_bottom_border(p_sk_hdr, theme_color_hex.replace("#", ""))
    p_loop_sk = doc.add_paragraph()
    _hidden_tag(p_loop_sk, "{% for group in skills %}", font_name)
    p_skill = doc.add_paragraph()
    run_sk = p_skill.add_run("{{ group.category }}: {{ group.skill_items | join(', ') }}")
    run_sk.font.name = font_name
    run_sk.font.size = Pt(10)
    p_loop_sk_end = doc.add_paragraph()
    _hidden_tag(p_loop_sk_end, "{% endfor %}", font_name)
    p_sk_end = doc.add_paragraph()
    _hidden_tag(p_sk_end, "{% endif %}", font_name)

    p_cert_if = doc.add_paragraph()
    _hidden_tag(p_cert_if, "{% if has_certifications %}", font_name)
    p_cert_hdr = doc.add_paragraph()
    run_cert_hdr = p_cert_hdr.add_run("CERTIFICATIONS")
    run_cert_hdr.font.name = font_name
    run_cert_hdr.font.size = Pt(13)
    run_cert_hdr.font.bold = True
    run_cert_hdr.font.color.rgb = theme_color
    p_loop_cert = doc.add_paragraph()
    _hidden_tag(p_loop_cert, "{% for item in certifications %}", font_name)
    p_cert = doc.add_paragraph(style="List Bullet")
    run_cert = p_cert.add_run("{{ item }}")
    run_cert.font.name = font_name
    run_cert.font.size = Pt(10)
    p_loop_cert_end = doc.add_paragraph()
    _hidden_tag(p_loop_cert_end, "{% endfor %}", font_name)
    p_cert_end = doc.add_paragraph()
    _hidden_tag(p_cert_end, "{% endif %}", font_name)

    p_ach_if = doc.add_paragraph()
    _hidden_tag(p_ach_if, "{% if has_achievements %}", font_name)
    p_ach_hdr = doc.add_paragraph()
    run_ach_hdr = p_ach_hdr.add_run("ACHIEVEMENTS")
    run_ach_hdr.font.name = font_name
    run_ach_hdr.font.size = Pt(13)
    run_ach_hdr.font.bold = True
    run_ach_hdr.font.color.rgb = theme_color
    p_loop_ach = doc.add_paragraph()
    _hidden_tag(p_loop_ach, "{% for item in achievements %}", font_name)
    p_ach = doc.add_paragraph(style="List Bullet")
    run_ach = p_ach.add_run("{{ item }}")
    run_ach.font.name = font_name
    run_ach.font.size = Pt(10)
    p_loop_ach_end = doc.add_paragraph()
    _hidden_tag(p_loop_ach_end, "{% endfor %}", font_name)
    p_ach_end = doc.add_paragraph()
    _hidden_tag(p_ach_end, "{% endif %}", font_name)

    output_path = TEMPLATE_DIR / filename
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Created template: {output_path}")


if __name__ == "__main__":
    create_template("canonical_minimalist.docx", "#475569", "Calibri")
    create_template("canonical_modern_blue.docx", "#1E3A8A", "Arial")
