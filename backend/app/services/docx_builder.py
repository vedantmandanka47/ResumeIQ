"""Build formatted DOCX files from rewritten resume text."""

import io
import re
from dataclasses import dataclass, field

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BODY_FONT = "Calibri"
BODY_SIZE = Pt(11)
HEADING_COLOR = "1F4E79"
COMMON_HEADINGS = {
    "summary",
    "professional summary",
    "profile",
    "experience",
    "work experience",
    "professional experience",
    "projects",
    "education",
    "skills",
    "technical skills",
    "certifications",
    "awards",
    "achievements",
}


@dataclass
class ResumeSection:
    title: str
    lines: list[str] = field(default_factory=list)


def _set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE

    heading = document.styles["Heading 1"]
    heading.font.name = BODY_FONT
    heading.font.size = Pt(14)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor.from_string(HEADING_COLOR)


def _add_bottom_border(paragraph, color: str = HEADING_COLOR) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def _clean_marker(line: str) -> str:
    return re.sub(r"^[-•*–]\s*", "", line).strip()


def _parse_sections(text: str) -> tuple[list[str], list[ResumeSection]]:
    intro: list[str] = []
    sections: list[ResumeSection] = []
    current: ResumeSection | None = None

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading_match = re.match(r"^#{1,3}\s+(.+)$", line)
        plain_heading = line.rstrip(":").strip()
        is_plain_heading = (
            plain_heading.lower() in COMMON_HEADINGS
            or (len(plain_heading) <= 48 and plain_heading.isupper())
        )
        if heading_match or is_plain_heading:
            title = heading_match.group(1).strip() if heading_match else plain_heading
            current = ResumeSection(title.rstrip(":"))
            sections.append(current)
            continue
        if current is None:
            intro.append(line)
        else:
            current.lines.append(line)

    if not sections and intro:
        sections.append(ResumeSection("Professional Summary", intro[1:] if len(intro) > 1 else intro))
        intro = intro[:1] if len(intro) > 1 else []
    return intro, sections


def _add_header(document: Document, intro: list[str]) -> None:
    if not intro:
        return

    name = intro[0]
    contact = " | ".join(intro[1:])

    name_paragraph = document.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_paragraph.add_run(name)
    name_run.bold = True
    name_run.font.name = BODY_FONT
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor.from_string("111827")

    if contact:
        contact_paragraph = document.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_paragraph.add_run(contact)
        contact_run.font.name = BODY_FONT
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor.from_string("4B5563")


def _add_section(document: Document, section: ResumeSection) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(4)
    heading.add_run(section.title.upper())
    _add_bottom_border(heading)

    for line in section.lines:
        is_bullet = line.startswith(("-", "•", "*", "–"))
        paragraph = document.add_paragraph(style="List Bullet" if is_bullet else None)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(_clean_marker(line) if is_bullet else line)
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE


def build_resume_docx(rewritten_text: str) -> bytes:
    """Return a professional DOCX document for markdown-light resume text."""
    document = Document()
    _set_document_defaults(document)

    intro, sections = _parse_sections(rewritten_text)
    _add_header(document, intro)

    for section in sections:
        if section.lines:
            _add_section(document, section)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
