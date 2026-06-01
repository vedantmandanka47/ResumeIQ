"""DOCX resume text extraction."""

import io
import logging

from docx import Document
from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

_COMMON_HEADINGS = {
    "summary",
    "professional summary",
    "profile",
    "experience",
    "work experience",
    "professional experience",
    "employment",
    "projects",
    "education",
    "skills",
    "technical skills",
    "certifications",
    "awards",
    "achievements",
    "publications",
    "leadership",
}


def _is_heading(paragraph: Paragraph, text: str) -> bool:
    style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
    normalized = text.rstrip(":").strip().lower()
    if "heading" in style_name or "title" in style_name:
        return True
    if normalized in _COMMON_HEADINGS:
        return True
    return len(text) <= 48 and text.isupper() and any(char.isalpha() for char in text)


def _is_bullet(paragraph: Paragraph, text: str) -> bool:
    style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
    if "list" in style_name or "bullet" in style_name:
        return True
    return text.startswith(("•", "-", "–", "*"))


def _annotate_paragraph(paragraph: Paragraph) -> str:
    text = " ".join(paragraph.text.split())
    if not text:
        return ""
    if _is_heading(paragraph, text):
        return f"## {text.rstrip(':')}"
    if _is_bullet(paragraph, text):
        return f"- {text.lstrip('•-*– ').strip()}"
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(file_bytes))
        blocks = [
            annotated
            for paragraph in document.paragraphs
            if (annotated := _annotate_paragraph(paragraph))
        ]

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        return "\n".join(blocks)
    except Exception as exc:
        logger.error("DOCX extraction failed: %s", exc, exc_info=True)
        return ""
