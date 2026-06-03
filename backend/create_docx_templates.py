import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bottom_border(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')  # Thickness of border
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_template(filename, theme_color_hex, font_name="Calibri"):
    doc = Document()
    
    # 1. Page Margins (0.75 inches for perfect balance)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    
    # Define styles
    theme_color = RGBColor.from_string(theme_color_hex.replace("#", ""))
    dark_gray = RGBColor.from_string("333333")
    light_gray = RGBColor.from_string("666666")
    
    # 2. Header Block (Name, Contact Details)
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    run_name = p_name.add_run("{{ name }}")
    run_name.font.name = font_name
    run_name.font.size = Pt(24)
    run_name.font.bold = True
    run_name.font.color.rgb = theme_color
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(14)
    
    contact_text = "{{ email }}  |  {{ phone }}  |  {{ location }}"
    run_contact = p_contact.add_run(contact_text)
    run_contact.font.name = font_name
    run_contact.font.size = Pt(10)
    run_contact.font.color.rgb = dark_gray
    
    # Dynamic Link section in contact if they exist
    # E.g. {% if linkedin %} | {{ linkedin }}{% endif %}{% if github %} | {{ github }}{% endif %}
    p_links = doc.add_paragraph()
    p_links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_links.paragraph_format.space_after = Pt(18)
    run_links = p_links.add_run("{% if linkedin %}{{ linkedin }}{% endif %}{% if github %}  |  {{ github }}{% endif %}{% if website %}  |  {{ website }}{% endif %}")
    run_links.font.name = font_name
    run_links.font.size = Pt(9.5)
    run_links.font.color.rgb = light_gray
    
    # 3. Professional Summary Section
    p_sum_hdr = doc.add_paragraph()
    p_sum_hdr.paragraph_format.space_before = Pt(12)
    p_sum_hdr.paragraph_format.space_after = Pt(4)
    p_sum_hdr.paragraph_format.keep_with_next = True
    run_sum_hdr = p_sum_hdr.add_run("PROFESSIONAL SUMMARY")
    run_sum_hdr.font.name = font_name
    run_sum_hdr.font.size = Pt(13)
    run_sum_hdr.font.bold = True
    run_sum_hdr.font.color.rgb = theme_color
    add_bottom_border(p_sum_hdr, theme_color_hex.replace("#", ""))
    
    p_sum_text = doc.add_paragraph()
    p_sum_text.paragraph_format.space_after = Pt(12)
    p_sum_text.paragraph_format.line_spacing = 1.15
    run_sum_text = p_sum_text.add_run("{{ summary }}")
    run_sum_text.font.name = font_name
    run_sum_text.font.size = Pt(10.5)
    run_sum_text.font.color.rgb = dark_gray
    
    # 4. Work Experience Section
    p_exp_hdr = doc.add_paragraph()
    p_exp_hdr.paragraph_format.space_before = Pt(14)
    p_exp_hdr.paragraph_format.space_after = Pt(6)
    p_exp_hdr.paragraph_format.keep_with_next = True
    run_exp_hdr = p_exp_hdr.add_run("PROFESSIONAL EXPERIENCE")
    run_exp_hdr.font.name = font_name
    run_exp_hdr.font.size = Pt(13)
    run_exp_hdr.font.bold = True
    run_exp_hdr.font.color.rgb = theme_color
    add_bottom_border(p_exp_hdr, theme_color_hex.replace("#", ""))
    
    # Loop start
    p_loop_exp_start = doc.add_paragraph()
    p_loop_exp_start.paragraph_format.space_before = Pt(0)
    p_loop_exp_start.paragraph_format.space_after = Pt(0)
    run_loop_exp_start = p_loop_exp_start.add_run("{% for exp in work_experience %}")
    run_loop_exp_start.font.name = font_name
    run_loop_exp_start.font.size = Pt(8)
    run_loop_exp_start.font.color.rgb = RGBColor.from_string("FFFFFF") # Hidden tag
    
    # Inside Loop: Job Title, Company, Dates, Location
    # We use a 2-column borderless table for alignment of Role/Company on left and Dates/Location on right
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    
    # Widths: Left col = 5.0 inches, Right col = 2.0 inches
    table.columns[0].width = Inches(5.0)
    table.columns[1].width = Inches(2.0)
    
    hdr_row = table.rows[0]
    cell_l = hdr_row.cells[0]
    cell_r = hdr_row.cells[1]
    
    set_cell_margins(cell_l, top=40, bottom=40, left=0, right=0)
    set_cell_margins(cell_r, top=40, bottom=40, left=0, right=0)
    
    p_l = cell_l.paragraphs[0]
    p_l.paragraph_format.space_after = Pt(2)
    run_role = p_l.add_run("{{ exp.role }}")
    run_role.font.name = font_name
    run_role.font.size = Pt(11)
    run_role.font.bold = True
    run_role.font.color.rgb = dark_gray
    
    run_sep = p_l.add_run("  |  ")
    run_sep.font.name = font_name
    run_sep.font.size = Pt(11)
    run_sep.font.color.rgb = light_gray
    
    run_comp = p_l.add_run("{{ exp.company }}")
    run_comp.font.name = font_name
    run_comp.font.size = Pt(11)
    run_comp.font.bold = True
    run_comp.font.color.rgb = theme_color
    
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_r.paragraph_format.space_after = Pt(2)
    run_dates = p_r.add_run("{{ exp.start_date }} - {{ exp.end_date }}")
    run_dates.font.name = font_name
    run_dates.font.size = Pt(10)
    run_dates.font.italic = True
    run_dates.font.color.rgb = dark_gray
    
    p_loc = doc.add_paragraph()
    p_loc.paragraph_format.space_before = Pt(0)
    p_loc.paragraph_format.space_after = Pt(4)
    run_loc = p_loc.add_run("{{ exp.location }}")
    run_loc.font.name = font_name
    run_loc.font.size = Pt(9.5)
    run_loc.font.italic = True
    run_loc.font.color.rgb = light_gray
    
    # Description Bullets Loop
    p_bullet_start = doc.add_paragraph()
    p_bullet_start.paragraph_format.space_before = Pt(0)
    p_bullet_start.paragraph_format.space_after = Pt(0)
    run_bullet_start = p_bullet_start.add_run("{% for bullet in exp.description %}")
    run_bullet_start.font.name = font_name
    run_bullet_start.font.size = Pt(8)
    run_bullet_start.font.color.rgb = RGBColor.from_string("FFFFFF") # Hidden tag
    
    # The actual bullet point
    p_bullet = doc.add_paragraph(style='List Bullet')
    p_bullet.paragraph_format.space_before = Pt(0)
    p_bullet.paragraph_format.space_after = Pt(2)
    p_bullet.paragraph_format.line_spacing = 1.1
    run_bullet = p_bullet.add_run("{{ bullet }}")
    run_bullet.font.name = font_name
    run_bullet.font.size = Pt(10)
    run_bullet.font.color.rgb = dark_gray
    
    p_bullet_end = doc.add_paragraph()
    p_bullet_end.paragraph_format.space_before = Pt(0)
    p_bullet_end.paragraph_format.space_after = Pt(4)
    run_bullet_end = p_bullet_end.add_run("{% endfor %}")
    run_bullet_end.font.name = font_name
    run_bullet_end.font.size = Pt(8)
    run_bullet_end.font.color.rgb = RGBColor.from_string("FFFFFF") # Hidden tag
    
    # End Experience Loop
    p_loop_exp_end = doc.add_paragraph()
    p_loop_exp_end.paragraph_format.space_before = Pt(0)
    p_loop_exp_end.paragraph_format.space_after = Pt(8)
    run_loop_exp_end = p_loop_exp_end.add_run("{% endfor %}")
    run_loop_exp_end.font.name = font_name
    run_loop_exp_end.font.size = Pt(8)
    run_loop_exp_end.font.color.rgb = RGBColor.from_string("FFFFFF") # Hidden tag
    
    # 5. Projects Section (if they exist)
    # Tag logic inside document using Jinja
    p_proj_start = doc.add_paragraph()
    p_proj_start.paragraph_format.space_before = Pt(0)
    p_proj_start.paragraph_format.space_after = Pt(0)
    run_proj_start = p_proj_start.add_run("{% if projects %}")
    run_proj_start.font.name = font_name
    run_proj_start.font.size = Pt(8)
    run_proj_start.font.color.rgb = RGBColor.from_string("FFFFFF")
    
    p_proj_hdr = doc.add_paragraph()
    p_proj_hdr.paragraph_format.space_before = Pt(12)
    p_proj_hdr.paragraph_format.space_after = Pt(6)
    p_proj_hdr.paragraph_format.keep_with_next = True
    run_proj_hdr = p_proj_hdr.add_run("PROJECTS")
    run_proj_hdr.font.name = font_name
    run_proj_hdr.font.size = Pt(13)
    run_proj_hdr.font.bold = True
    run_proj_hdr.font.color.rgb = theme_color
    add_bottom_border(p_proj_hdr, theme_color_hex.replace("#", ""))
    
    p_loop_proj_start = doc.add_paragraph()
    p_loop_proj_start.paragraph_format.space_before = Pt(0)
    p_loop_proj_start.paragraph_format.space_after = Pt(0)
    run_loop_proj_start = p_loop_proj_start.add_run("{% for proj in projects %}")
    run_loop_proj_start.font.name = font_name
    run_loop_proj_start.font.size = Pt(8)
    run_loop_proj_start.font.color.rgb = RGBColor.from_string("FFFFFF")
    
    # Inside Loop: Project Title, Tech Stack
    p_proj_title = doc.add_paragraph()
    p_proj_title.paragraph_format.space_before = Pt(4)
    p_proj_title.paragraph_format.space_after = Pt(2)
    run_proj_name = p_proj_title.add_run("{{ proj.name }}")
    run_proj_name.font.name = font_name
    run_proj_name.font.size = Pt(11)
    run_proj_name.font.bold = True
    run_proj_name.font.color.rgb = dark_gray
    
    run_proj_tech_sep = p_proj_title.add_run("  |  ")
    run_proj_tech_sep.font.name = font_name
    run_proj_tech_sep.font.size = Pt(10.5)
    run_proj_tech_sep.font.color.rgb = light_gray
    
    run_proj_tech = p_proj_title.add_run("{{ proj.technologies }}")
    run_proj_tech.font.name = font_name
    run_proj_tech.font.size = Pt(10.5)
    run_proj_tech.font.italic = True
    run_proj_tech.font.color.rgb = theme_color
    
    # Project bullets loop
    p_proj_bullet_start = doc.add_paragraph()
    p_proj_bullet_start.paragraph_format.space_before = Pt(0)
    p_proj_bullet_start.paragraph_format.space_after = Pt(0)
    run_proj_bullet_start = p_proj_bullet_start.add_run("{% for bullet in proj.description %}")
    run_proj_bullet_start.font.name = font_name
    run_proj_bullet_start.font.size = Pt(8)
    run_proj_bullet_start.font.color.rgb = RGBColor.from_string("FFFFFF")
    
    p_proj_bullet = doc.add_paragraph(style='List Bullet')
    p_proj_bullet.paragraph_format.space_before = Pt(0)
    p_proj_bullet.paragraph_format.space_after = Pt(2)
    run_proj_bullet_text = p_proj_bullet.add_run("{{ bullet }}")
    run_proj_bullet_text.font.name = font_name
    run_proj_bullet_text.font.size = Pt(10)
    run_proj_bullet_text.font.color.rgb = dark_gray
    
    p_proj_bullet_end = doc.add_paragraph()
    p_proj_bullet_end.paragraph_format.space_before = Pt(0)
    p_proj_bullet_end.paragraph_format.space_after = Pt(4)
    run_proj_bullet_end = p_proj_bullet_end.add_run("{% endfor %}")
    run_proj_bullet_end.font.name = font_name
    run_proj_bullet_end.font.size = Pt(8)
    run_proj_bullet_end.font.color.rgb = RGBColor.from_string("FFFFFF")
    
    p_loop_proj_end = doc.add_paragraph()
    p_loop_proj_end.paragraph_format.space_before = Pt(0)
    p_loop_proj_end.paragraph_format.space_after = Pt(6)
    run_loop_proj_end = p_loop_proj_end.add_run("{% endfor %}")
    run_loop_proj_end.font.name = font_name
    run_loop_proj_end.font.size = Pt(8)
    run_loop_proj_end.font.color.rgb = RGBColor.from_string("FFFFFF")
    
    p_proj_end = doc.add_paragraph()
    p_proj_end.paragraph_format.space_before = Pt(0)
    p_proj_end.paragraph_format.space_after = Pt(4)
    run_proj_end = p_proj_end.add_run("{% endif %}")
    run_proj_end.font.name = font_name
    run_proj_end.font.size = Pt(8)
    run_proj_end.font.color.rgb = RGBColor.from_string("FFFFFF")
    
    # 6. Education Section
    p_edu_hdr = doc.add_paragraph()
    p_edu_hdr.paragraph_format.space_before = Pt(12)
    p_edu_hdr.paragraph_format.space_after = Pt(6)
    p_edu_hdr.paragraph_format.keep_with_next = True
    run_edu_hdr = p_edu_hdr.add_run("EDUCATION")
    run_edu_hdr.font.name = font_name
    run_edu_hdr.font.size = Pt(13)
    run_edu_hdr.font.bold = True
    run_edu_hdr.font.color.rgb = theme_color
    add_bottom_border(p_edu_hdr, theme_color_hex.replace("#", ""))
    
    p_loop_edu_start = doc.add_paragraph()
    p_loop_edu_start.paragraph_format.space_before = Pt(0)
    p_loop_edu_start.paragraph_format.space_after = Pt(0)
    run_loop_edu_start = p_loop_edu_start.add_run("{% for edu in education %}")
    run_loop_edu_start.font.name = font_name
    run_loop_edu_start.font.size = Pt(8)
    run_loop_edu_start.font.color.rgb = RGBColor.from_string("FFFFFF")
    
    # Education 2-column table
    edu_table = doc.add_table(rows=1, cols=2)
    edu_table.autofit = False
    edu_table.columns[0].width = Inches(5.0)
    edu_table.columns[1].width = Inches(2.0)
    
    edu_row = edu_table.rows[0]
    edu_l = edu_row.cells[0]
    edu_r = edu_row.cells[1]
    
    set_cell_margins(edu_l, top=40, bottom=40, left=0, right=0)
    set_cell_margins(edu_r, top=40, bottom=40, left=0, right=0)
    
    p_edu_l = edu_l.paragraphs[0]
    p_edu_l.paragraph_format.space_after = Pt(2)
    run_degree = p_edu_l.add_run("{{ edu.degree }}")
    run_degree.font.name = font_name
    run_degree.font.size = Pt(11)
    run_degree.font.bold = True
    run_degree.font.color.rgb = dark_gray
    
    run_edu_sep = p_edu_l.add_run("  |  ")
    run_edu_sep.font.name = font_name
    run_edu_sep.font.size = Pt(11)
    run_edu_sep.font.color.rgb = light_gray
    
    run_school = p_edu_l.add_run("{{ edu.school }}")
    run_school.font.name = font_name
    run_school.font.size = Pt(11)
    run_school.font.bold = True
    run_school.font.color.rgb = theme_color
    
    p_edu_r = edu_r.paragraphs[0]
    p_edu_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_edu_r.paragraph_format.space_after = Pt(2)
    run_grad = p_edu_r.add_run("{{ edu.graduation_date }}")
    run_grad.font.name = font_name
    run_grad.font.size = Pt(10)
    run_grad.font.italic = True
    run_grad.font.color.rgb = dark_gray
    
    p_edu_sub = doc.add_paragraph()
    p_edu_sub.paragraph_format.space_before = Pt(0)
    p_edu_sub.paragraph_format.space_after = Pt(6)
    run_edu_loc = p_edu_sub.add_run("{{ edu.location }}")
    run_edu_loc.font.name = font_name
    run_edu_loc.font.size = Pt(9.5)
    run_edu_loc.font.italic = True
    run_edu_loc.font.color.rgb = light_gray
    
    run_gpa_check = p_edu_sub.add_run("{% if edu.gpa %}  |  GPA: {{ edu.gpa }}{% endif %}")
    run_gpa_check.font.name = font_name
    run_gpa_check.font.size = Pt(9.5)
    run_gpa_check.font.color.rgb = dark_gray
    
    p_loop_edu_end = doc.add_paragraph()
    p_loop_edu_end.paragraph_format.space_before = Pt(0)
    p_loop_edu_end.paragraph_format.space_after = Pt(6)
    run_loop_edu_end = p_loop_edu_end.add_run("{% endfor %}")
    run_loop_edu_end.font.name = font_name
    run_loop_edu_end.font.size = Pt(8)
    run_loop_edu_end.font.color.rgb = RGBColor.from_string("FFFFFF")
    
    # 7. Skills Section
    p_sk_hdr = doc.add_paragraph()
    p_sk_hdr.paragraph_format.space_before = Pt(12)
    p_sk_hdr.paragraph_format.space_after = Pt(6)
    p_sk_hdr.paragraph_format.keep_with_next = True
    run_sk_hdr = p_sk_hdr.add_run("SKILLS & COMPETENCIES")
    run_sk_hdr.font.name = font_name
    run_sk_hdr.font.size = Pt(13)
    run_sk_hdr.font.bold = True
    run_sk_hdr.font.color.rgb = theme_color
    add_bottom_border(p_sk_hdr, theme_color_hex.replace("#", ""))
    
    p_loop_sk_start = doc.add_paragraph()
    p_loop_sk_start.paragraph_format.space_before = Pt(0)
    p_loop_sk_start.paragraph_format.space_after = Pt(0)
    run_loop_sk_start = p_loop_sk_start.add_run("{% for skill in skills %}")
    run_loop_sk_start.font.name = font_name
    run_loop_sk_start.font.size = Pt(8)
    run_loop_sk_start.font.color.rgb = RGBColor.from_string("FFFFFF")
    
    # Key-value pair skill lines
    p_skill_line = doc.add_paragraph()
    p_skill_line.paragraph_format.space_before = Pt(0)
    p_skill_line.paragraph_format.space_after = Pt(3)
    p_skill_line.paragraph_format.left_indent = Inches(0.15)
    
    run_sk_cat = p_skill_line.add_run("{{ skill.category }}: ")
    run_sk_cat.font.name = font_name
    run_sk_cat.font.size = Pt(10)
    run_sk_cat.font.bold = True
    run_sk_cat.font.color.rgb = dark_gray
    
    run_sk_items = p_skill_line.add_run("{{ skill.items | join(', ') }}")
    run_sk_items.font.name = font_name
    run_sk_items.font.size = Pt(10)
    run_sk_items.font.color.rgb = dark_gray
    
    p_loop_sk_end = doc.add_paragraph()
    p_loop_sk_end.paragraph_format.space_before = Pt(0)
    p_loop_sk_end.paragraph_format.space_after = Pt(6)
    run_loop_sk_end = p_loop_sk_end.add_run("{% endfor %}")
    run_loop_sk_end.font.name = font_name
    run_loop_sk_end.font.size = Pt(8)
    run_loop_sk_end.font.color.rgb = RGBColor.from_string("FFFFFF")
    
    # Save the file
    os.makedirs(os.path.dirname(filename), exist_ok=True)
    doc.save(filename)
    print(f"Created template: {filename}")

if __name__ == "__main__":
    # Create the two templates with elegant color schemes!
    # 1. Minimalist Slate
    create_template(r"c:\Vedant\Self Drive A\Project\Deppost\ResumeIQ\backend\templete\Minimalist.docx", "#475569", "Calibri")
    # 2. Modern Navy
    create_template(r"c:\Vedant\Self Drive A\Project\Deppost\ResumeIQ\backend\templete\Minimalist-1.docx", "#1E3A8A", "Arial")
